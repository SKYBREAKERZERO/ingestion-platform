from __future__ import annotations

import re
from pathlib import Path
from typing import Iterator

from docx import Document as WordDocument
from docx.document import Document as WordDocumentType
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.loader.base_loader import BaseLoader
from app.model.block import (
    BlockType,
    DocumentBlock,
)
from app.model.document import Document
from app.model.page import Page


class DOCXLoaderError(RuntimeError):
    """
    DOCX 文件加载异常。
    """


class DOCXLoader(BaseLoader):
    """
    DOCX 企业级原始结构加载器。

    负责：
        - 按 Word document.xml 原始顺序读取 Paragraph / Table
        - 尽量无损保留表格 Cell 列位置
        - 保留重复 Cell
        - 保留 0 / False 等文本化后的有效值
        - 保留 Word Style 信息
        - 多信号识别 Heading 层级
        - 保存 Heading 层级来源与冲突信息
        - 输出统一 Document.pages
        - 输出结构化 Document.blocks
        - 保存 Loader Metadata

    Heading 层级识别来源：
        1. Paragraph outlineLvl
        2. Style outlineLvl
        3. Heading style name / style id
        4. Numbering ilvl 作为 Heading 的辅助信号

    不负责：
        - Chapter / Section 建模
        - Paragraph 内容过滤
        - Table 内容过滤
        - 标题合并
        - Chunk
        - Token 统计
        - JSON / PostgreSQL 保存

    设计原则：
        Loader 只负责“读取”和“保存来源信息”，
        不提前执行不可逆的业务语义删除。

    关于 page_number：
        python-docx 无法可靠提供 Word 最终排版后的物理页码。

        因此：
            DocumentBlock.page_number = None

        Document.pages 仍保留一个逻辑 Page(page_number=1)，
        仅用于兼容现有统一 Document 模型。

        真实物理页码不可用时，宁可返回 None，
        不把所有正文错误标记成第 1 页。
    """

    # ==================================================
    # Heading Style
    # ==================================================

    _HEADING_STYLE_NAME_PATTERN = re.compile(
        r"^(?:heading|見出し|标题|標題)\s*([1-9]\d*)$",
        re.IGNORECASE,
    )

    _HEADING_STYLE_ID_PATTERN = re.compile(
        r"^heading([1-9]\d*)$",
        re.IGNORECASE,
    )

    # Word outlineLvl / numPr ilvl 是 0-based。
    _MIN_HEADING_LEVEL = 1

    # 防御异常文档，避免极端 outlineLvl 污染后续结构。
    _MAX_HEADING_LEVEL = 9

    # ==================================================
    # Public API
    # ==================================================

    def load(
        self,
        file_path: str,
    ) -> Document:
        """
        加载 DOCX 文件。

        Args:
            file_path:
                DOCX 文件路径。

        Returns:
            原始结构化 Document。
        """

        path = self._validate_path(
            file_path
        )

        try:

            word_document = WordDocument(
                str(path)
            )

            return self._build_document(
                path=path,
                word_document=word_document,
            )

        except DOCXLoaderError:
            raise

        except Exception as exc:

            raise DOCXLoaderError(
                "Failed to load DOCX file "
                f"'{path.name}': {exc}"
            ) from exc

    # ==================================================
    # Build Document
    # ==================================================

    def _build_document(
        self,
        *,
        path: Path,
        word_document: WordDocumentType,
    ) -> Document:

        blocks: list[
            DocumentBlock
        ] = []

        plain_text_lines: list[
            str
        ] = []

        paragraph_count = 0
        heading_count = 0
        list_count = 0

        table_count = 0
        table_row_count = 0

        empty_table_cell_count = 0
        table_cell_count = 0

        style_heading_count = 0
        paragraph_outline_heading_count = 0
        style_outline_heading_count = 0
        numbering_assisted_heading_count = 0
        heading_level_conflict_count = 0

        order = 0

        # ==============================================
        # Original Word Order
        # ==============================================

        for raw_block in self._iter_blocks(
            word_document
        ):

            # ==========================================
            # Paragraph
            # ==========================================

            if isinstance(
                raw_block,
                Paragraph,
            ):

                text = self._normalize_text(
                    raw_block.text
                )

                # 空 Paragraph 不建立 Block。
                #
                # 这里仅删除完全无文本内容的排版段落，
                # 不删除任何有实际文本的 Paragraph。
                if not text:
                    continue

                style_name = (
                    self._get_style_name(
                        raw_block
                    )
                )

                style_id = (
                    self._get_style_id(
                        raw_block
                    )
                )

                # ======================================
                # Heading Signals
                # ======================================

                style_heading_level = (
                    self._get_style_heading_level(
                        style_name=style_name,
                        style_id=style_id,
                    )
                )

                paragraph_outline_level = (
                    self._get_paragraph_outline_heading_level(
                        raw_block
                    )
                )

                style_outline_level = (
                    self._get_style_outline_heading_level(
                        raw_block
                    )
                )

                numbering_level = (
                    self._get_numbering_level(
                        raw_block
                    )
                )

                numbering_id = (
                    self._get_numbering_id(
                        raw_block
                    )
                )

                (
                    heading_level,
                    heading_level_source,
                    heading_level_conflict,
                ) = self._resolve_heading_level(
                    style_heading_level=(
                        style_heading_level
                    ),
                    paragraph_outline_level=(
                        paragraph_outline_level
                    ),
                    style_outline_level=(
                        style_outline_level
                    ),
                    numbering_level=(
                        numbering_level
                    ),
                )

                if heading_level_conflict:
                    heading_level_conflict_count += 1

                # ======================================
                # Block Type
                # ======================================

                if heading_level is not None:

                    block_type = (
                        BlockType.HEADING
                    )

                    heading_count += 1

                    if (
                        heading_level_source
                        == "paragraph_outline"
                    ):

                        paragraph_outline_heading_count += 1

                    elif (
                        heading_level_source
                        == "style_outline"
                    ):

                        style_outline_heading_count += 1

                    elif (
                        heading_level_source
                        == "style"
                    ):

                        style_heading_count += 1

                    elif (
                        heading_level_source
                        == "numbering_assisted"
                    ):

                        numbering_assisted_heading_count += 1

                elif self._is_list_paragraph(
                    raw_block
                ):

                    block_type = (
                        BlockType.LIST
                    )

                    list_count += 1

                else:

                    block_type = (
                        BlockType.PARAGRAPH
                    )

                # ======================================
                # Paragraph Block
                # ======================================

                block = DocumentBlock(
                    block_type=block_type,
                    text=text,
                    level=heading_level,
                    style_name=style_name,
                    order=order,

                    # DOCX 无可靠物理页码。
                    page_number=None,

                    source="docx",

                    metadata={
                        "source": (
                            "paragraph"
                        ),

                        "style_name": (
                            style_name
                        ),

                        "style_id": (
                            style_id
                        ),

                        "style_heading_level": (
                            style_heading_level
                        ),

                        "paragraph_outline_heading_level": (
                            paragraph_outline_level
                        ),

                        "style_outline_heading_level": (
                            style_outline_level
                        ),

                        "numbering_level": (
                            numbering_level
                        ),

                        "numbering_id": (
                            numbering_id
                        ),

                        "heading_level_source": (
                            heading_level_source
                        ),

                        "heading_level_conflict": (
                            heading_level_conflict
                        ),

                        "physical_page_number_available": (
                            False
                        ),
                    },
                )

                blocks.append(
                    block
                )

                plain_text_lines.append(
                    text
                )

                paragraph_count += 1
                order += 1

                continue

            # ==========================================
            # Table
            # ==========================================

            if isinstance(
                raw_block,
                Table,
            ):

                current_table_index = (
                    table_count
                )

                table_count += 1

                for (
                    row_index,
                    row,
                ) in enumerate(
                    raw_block.rows
                ):

                    # ==================================
                    # Extract Cells
                    # ==================================
                    #
                    # 不删除空 Cell。
                    # 不删除重复 Cell。
                    #
                    # Loader 必须保留原始列位置。

                    values = (
                        self._extract_row_values(
                            row.cells
                        )
                    )

                    table_cell_count += len(
                        values
                    )

                    empty_table_cell_count += sum(
                        1
                        for value
                        in values
                        if not value
                    )

                    # ==================================
                    # Completely Empty Row
                    # ==================================

                    if not any(
                        values
                    ):
                        continue

                    # ==================================
                    # Display Text
                    # ==================================

                    row_text = (
                        self._build_row_text(
                            values
                        )
                    )

                    # ==================================
                    # Table Block
                    # ==================================

                    block = DocumentBlock(
                        block_type=BlockType.TABLE,

                        text=row_text,

                        order=order,

                        table_index=(
                            current_table_index
                        ),

                        row_index=row_index,

                        cells=values,

                        # DOCX 无可靠物理页码。
                        page_number=None,

                        source="docx",

                        metadata={
                            "source": (
                                "table"
                            ),

                            "table_index": (
                                current_table_index
                            ),

                            "row_index": (
                                row_index
                            ),

                            "column_count": (
                                len(values)
                            ),

                            "non_empty_cell_count": (
                                sum(
                                    1
                                    for value
                                    in values
                                    if value
                                )
                            ),

                            "physical_page_number_available": (
                                False
                            ),
                        },
                    )

                    blocks.append(
                        block
                    )

                    plain_text_lines.append(
                        row_text
                    )

                    table_row_count += 1
                    order += 1

        # ==============================================
        # Logical Page
        # ==============================================
        #
        # 这里的 Page 1 仅是统一模型中的“逻辑 Page”。
        #
        # 不表示 Word 文档真实物理第 1 页。

        content = "\n".join(
            plain_text_lines
        ).strip()

        # ==============================================
        # Metadata
        # ==============================================

        metadata = {
            "source_format": (
                "docx"
            ),

            "loader": (
                "DOCXLoader"
            ),

            "loader_status": (
                "SUCCESS"
            ),

            "paragraph_count": (
                paragraph_count
            ),

            "heading_count": (
                heading_count
            ),

            "list_count": (
                list_count
            ),

            "table_count": (
                table_count
            ),

            "table_row_count": (
                table_row_count
            ),

            "table_cell_count": (
                table_cell_count
            ),

            "empty_table_cell_count": (
                empty_table_cell_count
            ),

            "block_count": (
                len(blocks)
            ),

            "character_count": (
                len(content)
            ),

            # Heading signal diagnostics
            "style_heading_count": (
                style_heading_count
            ),

            "paragraph_outline_heading_count": (
                paragraph_outline_heading_count
            ),

            "style_outline_heading_count": (
                style_outline_heading_count
            ),

            "numbering_assisted_heading_count": (
                numbering_assisted_heading_count
            ),

            "heading_level_conflict_count": (
                heading_level_conflict_count
            ),

            # DOCX physical pagination diagnostics
            "page_number_strategy": (
                "unavailable_for_docx_layout"
            ),

            "physical_page_number_available": (
                False
            ),

            "logical_page_count": (
                1
            ),
        }

        return Document(
            file_name=path.name,

            file_type="docx",

            pages=[
                Page(
                    page_number=1,
                    text=content,
                )
            ],

            blocks=blocks,

            metadata=metadata,
        )

    # ==================================================
    # Heading Level Resolution
    # ==================================================

    @classmethod
    def _resolve_heading_level(
        cls,
        *,
        style_heading_level: int | None,
        paragraph_outline_level: int | None,
        style_outline_level: int | None,
        numbering_level: int | None,
    ) -> tuple[
        int | None,
        str | None,
        bool,
    ]:
        """
        综合多个 Word OOXML 信号解析 Heading Level。

        可信度顺序：
            1. Paragraph outlineLvl
            2. Style outlineLvl
            3. Heading style
            4. numbering ilvl 辅助修正

        numbering ilvl 本身不会把普通 List 强行升级成 Heading。

        但如果 Paragraph 已经由 Style / outlineLvl
        确认为 Heading，而 numbering ilvl 与 Heading 层级冲突，
        则 numbering 可用于恢复“所有标题错误设置为 Heading 1”
        这类企业文档常见格式问题。
        """

        primary_candidates: list[
            tuple[
                str,
                int | None,
            ]
        ] = [
            (
                "paragraph_outline",
                paragraph_outline_level,
            ),
            (
                "style_outline",
                style_outline_level,
            ),
            (
                "style",
                style_heading_level,
            ),
        ]

        existing = [
            (
                source,
                cls._clamp_heading_level(
                    level
                ),
            )
            for source, level
            in primary_candidates
            if level is not None
        ]

        if not existing:
            return (
                None,
                None,
                False,
            )

        existing_levels = {
            level
            for _, level
            in existing
        }

        conflict = (
            len(existing_levels)
            > 1
        )

        source, resolved_level = (
            existing[0]
        )

        # ==============================================
        # Numbering-assisted hierarchy recovery
        # ==============================================

        if numbering_level is not None:

            numbering_heading_level = (
                cls._clamp_heading_level(
                    numbering_level + 1
                )
            )

            if (
                numbering_heading_level
                != resolved_level
            ):

                resolved_level = (
                    numbering_heading_level
                )

                source = (
                    "numbering_assisted"
                )

                conflict = True

        return (
            resolved_level,
            source,
            conflict,
        )

    @classmethod
    def _clamp_heading_level(
        cls,
        level: int,
    ) -> int:

        normalized = int(
            level
        )

        return max(
            cls._MIN_HEADING_LEVEL,
            min(
                normalized,
                cls._MAX_HEADING_LEVEL,
            ),
        )

    # ==================================================
    # Heading Style
    # ==================================================

    @classmethod
    def _get_style_heading_level(
        cls,
        *,
        style_name: str | None,
        style_id: str | None,
    ) -> int | None:

        for (
            value,
            pattern,
        ) in (
            (
                style_name,
                cls._HEADING_STYLE_NAME_PATTERN,
            ),
            (
                style_id,
                cls._HEADING_STYLE_ID_PATTERN,
            ),
        ):

            if not value:
                continue

            normalized = (
                cls._normalize_text(
                    value
                )
            )

            match = pattern.fullmatch(
                normalized
            )

            if match is None:
                continue

            return cls._clamp_heading_level(
                int(
                    match.group(1)
                )
            )

        return None

    # ==================================================
    # Paragraph Outline Level
    # ==================================================

    @classmethod
    def _get_paragraph_outline_heading_level(
        cls,
        paragraph: Paragraph,
    ) -> int | None:
        """
        读取 Paragraph 自身：

            w:pPr/w:outlineLvl

        OOXML outlineLvl:
            0 -> Heading level 1
            1 -> Heading level 2
            ...
        """

        paragraph_element = (
            paragraph._p
        )

        try:

            outline_nodes = (
                paragraph_element.xpath(
                    "./w:pPr/w:outlineLvl"
                )
            )

        except Exception:

            return None

        if not outline_nodes:
            return None

        value = outline_nodes[
            0
        ].get(
            qn(
                "w:val"
            )
        )

        return cls._convert_zero_based_level(
            value
        )

    # ==================================================
    # Style Outline Level
    # ==================================================

    @classmethod
    def _get_style_outline_heading_level(
        cls,
        paragraph: Paragraph,
    ) -> int | None:
        """
        读取 Paragraph 所属 Style：

            w:style/w:pPr/w:outlineLvl
        """

        style = (
            paragraph.style
        )

        if style is None:
            return None

        style_element = getattr(
            style,
            "element",
            None,
        )

        if style_element is None:
            return None

        try:

            outline_nodes = (
                style_element.xpath(
                    "./w:pPr/w:outlineLvl"
                )
            )

        except Exception:

            return None

        if not outline_nodes:
            return None

        value = outline_nodes[
            0
        ].get(
            qn(
                "w:val"
            )
        )

        return cls._convert_zero_based_level(
            value
        )

    # ==================================================
    # Numbering
    # ==================================================

    @staticmethod
    def _get_numbering_level(
        paragraph: Paragraph,
    ) -> int | None:
        """
        返回 Word numbering ilvl。

        注意：
            ilvl 是 0-based。
        """

        properties = (
            paragraph._p.pPr
        )

        if (
            properties is None
            or properties.numPr is None
        ):
            return None

        ilvl = (
            properties.numPr.ilvl
        )

        if ilvl is None:
            return None

        try:

            return int(
                ilvl.val
            )

        except (
            TypeError,
            ValueError,
        ):

            return None

    @staticmethod
    def _get_numbering_id(
        paragraph: Paragraph,
    ) -> int | None:

        properties = (
            paragraph._p.pPr
        )

        if (
            properties is None
            or properties.numPr is None
        ):
            return None

        num_id = (
            properties.numPr.numId
        )

        if num_id is None:
            return None

        try:

            return int(
                num_id.val
            )

        except (
            TypeError,
            ValueError,
        ):

            return None

    # ==================================================
    # Style Metadata
    # ==================================================

    @staticmethod
    def _get_style_name(
        paragraph: Paragraph,
    ) -> str | None:

        style = (
            paragraph.style
        )

        if style is None:
            return None

        name = getattr(
            style,
            "name",
            None,
        )

        if not name:
            return None

        normalized = str(
            name
        ).strip()

        return (
            normalized
            or None
        )

    @staticmethod
    def _get_style_id(
        paragraph: Paragraph,
    ) -> str | None:

        style = (
            paragraph.style
        )

        if style is None:
            return None

        style_id = getattr(
            style,
            "style_id",
            None,
        )

        if not style_id:
            return None

        normalized = str(
            style_id
        ).strip()

        return (
            normalized
            or None
        )

    # ==================================================
    # OOXML Level
    # ==================================================

    @classmethod
    def _convert_zero_based_level(
        cls,
        value,
    ) -> int | None:

        if value is None:
            return None

        try:

            zero_based_level = int(
                value
            )

        except (
            TypeError,
            ValueError,
        ):

            return None

        if zero_based_level < 0:
            return None

        return cls._clamp_heading_level(
            zero_based_level + 1
        )

    # ==================================================
    # List Detection
    # ==================================================

    @staticmethod
    def _is_list_paragraph(
        paragraph: Paragraph,
    ) -> bool:
        """
        判断 Paragraph 是否包含 Word 编号或项目符号属性。
        """

        properties = (
            paragraph._p.pPr
        )

        if properties is None:
            return False

        return (
            properties.numPr
            is not None
        )

    # ==================================================
    # Path Validation
    # ==================================================

    @staticmethod
    def _validate_path(
        file_path: str,
    ) -> Path:

        if not file_path:

            raise ValueError(
                "file_path cannot be empty."
            )

        path = Path(
            file_path
        ).expanduser()

        if not path.exists():

            raise FileNotFoundError(
                "DOCX file not found: "
                f"{path}"
            )

        if not path.is_file():

            raise IsADirectoryError(
                "Input path is not a file: "
                f"{path}"
            )

        # Word 打开的临时文件。
        if path.name.startswith(
            "~$"
        ):

            raise ValueError(
                "Temporary Word file "
                "is not supported: "
                f"{path.name}"
            )

        if (
            path.suffix.lower()
            != ".docx"
        ):

            raise ValueError(
                "DOCXLoader only accepts "
                ".docx files. "
                f"Received: "
                f"{path.suffix or '<no extension>'}"
            )

        return path

    # ==================================================
    # Normalize Text
    # ==================================================

    @staticmethod
    def _normalize_text(
        text: str | None,
    ) -> str:
        """
        Loader 级别轻量文本标准化。

        不执行具有业务语义的改写。
        """

        if text is None:
            return ""

        normalized = str(
            text
        )

        normalized = (
            normalized.replace(
                "\u3000",
                " ",
            )
        )

        normalized = (
            normalized.replace(
                "\xa0",
                " ",
            )
        )

        normalized = (
            normalized.replace(
                "\r\n",
                "\n",
            )
        )

        normalized = (
            normalized.replace(
                "\r",
                "\n",
            )
        )

        lines = [
            " ".join(
                line.split()
            )
            for line
            in normalized.splitlines()
            if line.strip()
        ]

        # 对同一个 Paragraph / Cell 内的多行，
        # 保留显式分隔，避免直接粘连。
        return " / ".join(
            lines
        ).strip()

    # ==================================================
    # Extract Table Row
    # ==================================================

    @classmethod
    def _extract_row_values(
        cls,
        cells,
    ) -> list[str]:
        """
        提取一行所有 Cell。

        保留：
            - 空 Cell
            - 重复 Cell
            - 原始列位置
        """

        values: list[
            str
        ] = []

        for cell in cells:

            value = (
                cls._normalize_text(
                    cell.text
                )
            )

            values.append(
                value
            )

        return values

    # ==================================================
    # Build Row Text
    # ==================================================

    @staticmethod
    def _build_row_text(
        values: list[str],
    ) -> str:
        """
        构建用于 Page / Block 的逻辑表格文本。

        cells 本身仍完整保存列位置。
        """

        return " | ".join(
            values
        ).strip()

    # ==================================================
    # Iterate Word Blocks
    # ==================================================

    @staticmethod
    def _iter_blocks(
        document: WordDocumentType,
    ) -> Iterator[
        Paragraph | Table
    ]:
        """
        按 document.xml 原始顺序遍历：

            Paragraph
            Table

        避免分别读取：

            document.paragraphs
            document.tables

        导致 Paragraph / Table 相对顺序丢失。
        """

        parent_element = (
            document.element.body
        )

        for child in (
            parent_element.iterchildren()
        ):

            if isinstance(
                child,
                CT_P,
            ):

                yield Paragraph(
                    child,
                    document,
                )

            elif isinstance(
                child,
                CT_Tbl,
            ):

                yield Table(
                    child,
                    document,
                )
